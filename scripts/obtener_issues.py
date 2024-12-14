#python -m venv env
#.\env\Scripts\activate
#pip install requests python-docx
#pip install python-dotenv
#
#pip freeze > requirements.txt
#pip install -r requirements.txt
#
#
# python .\obtener_issues.py
#
import requests
from datetime import datetime
from docx import Document
from dotenv import load_dotenv
import os

load_dotenv()

# Configuración
GITHUB_API_URL = "https://api.github.com"
REPO_OWNER = "UPT-FAING-EPIS"
REPO_NAME = "proyecto-si8811a-2024-ii-u2-api-y-funciones-anahua"
TOKEN = os.getenv("GITHUB_TOKEN")

def obtener_issues():
    url = f"{GITHUB_API_URL}/repos/{REPO_OWNER}/{REPO_NAME}/issues"
    headers = {"Authorization": f"token {TOKEN}"}
    params = {"state": "all"}

    # Obtener el directorio donde está el script
    current_dir = os.path.dirname(os.path.abspath(__file__))

    # Crear carpeta 'obtener_issues' si no existe
    output_dir = os.path.join(current_dir, "obtener_issues")
    os.makedirs(output_dir, exist_ok=True)

    response = requests.get(url, headers=headers, params=params)

    if response.status_code == 200:
        issues = response.json()

        # Generar el archivo .md (en formato tabla)
        archivo_md = os.path.join(output_dir, "issues_informe.md")
        with open(archivo_md, "w", encoding="utf-8") as file:
            file.write(f"# Issues del Repositorio '{REPO_NAME}'\n\n")
            file.write("Tabla 1: Issues del repositorio\n\n")
            file.write("| Número | Título                              | Estado  | Creado por | URL                                       |\n")
            file.write("|--------|------------------------------------|---------|------------|-------------------------------------------|\n")
            for issue in issues:
                file.write(f"| #{issue['number']} | {issue['title'][:30]:<30} | {issue['state']} | {issue['user']['login']} | [Enlace]({issue['html_url']}) |\n")
            file.write("\nFuente: Elaboración propia\n")
        print(f"Archivo '{archivo_md}' generado exitosamente.")

        # Generar el archivo .docx (en formato tabla)
        archivo_docx = os.path.join(output_dir, "issues_informe.docx")
        doc = Document()
        doc.add_heading(f"Issues del Repositorio '{REPO_NAME}'", level=1)

        doc.add_paragraph("Tabla 1: Issues del repositorio", style="Title").alignment = 1
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"

        # Encabezados de la tabla
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Número"
        hdr_cells[1].text = "Título"
        hdr_cells[2].text = "Estado"
        hdr_cells[3].text = "Creado por"
        hdr_cells[4].text = "URL"

        # Llenar la tabla con datos
        for issue in issues:
            row_cells = table.add_row().cells
            row_cells[0].text = f"#{issue['number']}"
            row_cells[1].text = issue['title']
            row_cells[2].text = issue['state']
            row_cells[3].text = issue['user']['login']
            row_cells[4].text = issue['html_url']

        doc.add_paragraph("\nFuente: Elaboración propia", style="Normal").alignment = 1
        doc.save(archivo_docx)
        print(f"Archivo '{archivo_docx}' generado exitosamente.")
    else:
        print(f"Error al obtener los issues: {response.status_code} - {response.text}")

if __name__ == "__main__":
    obtener_issues()
